import xml.etree.ElementTree as ET
import os
from docx import Document
from docx.shared import Inches

# Added to make data types readable
SSIS_TYPE_MAP = {
    "11": "DT_BOOL (Boolean)",
    "2": "DT_I2 (Int16)",
    "3": "DT_I4 (Int32)",
    "20": "DT_I8 (Int64)",
    "4": "DT_R4 (Float)",
    "5": "DT_R8 (Double)",
    "6": "DT_CY (Currency)",
    "7": "DT_DATE (Date)",
    "128": "DT_BYTES (Binary)",
    "129": "DT_STR (Ansi String)",
    "130": "DT_WSTR (Unicode String)",
    "131": "DT_NUMERIC (Decimal)",
    "133": "DT_DBDATE (Database Date)",
    "134": "DT_DBTIME (Database Time)",
    "135": "DT_DBTIMESTAMP (Database Timestamp)",
    "141": "DT_NUMERIC (High Precision)",
    "145": "DT_DBTIME2 (Database Time 2)",
    "302": "DT_TEXT (Long Ansi String)",
    "303": "DT_NTEXT (Long Unicode String)"
}

class DTSXParser:
    def __init__(self, file_path):
        self.file_path = file_path
        self.tree = ET.parse(file_path)
        self.root = self.tree.getroot()
        self.ns = {'DTS': 'www.microsoft.com/SqlServer/Dts'}
        self.connections = []
        self.variables = []
        self.sql_queries = []
        self.flat_file_columns = []

    def parse(self):
        self._parse_connections()
        self._parse_variables()
        self._parse_sql_and_flat_files()

    def _parse_connections(self):
        conn_managers = self.root.find('DTS:ConnectionManagers', self.ns)
        if conn_managers is not None:
            for conn in conn_managers.findall('DTS:ConnectionManager', self.ns):
                name = conn.get(f'{{{self.ns["DTS"]}}}ObjectName')
                creation_name = conn.get(f'{{{self.ns["DTS"]}}}CreationName')
                
                # Look for connection string in ObjectData
                conn_str = ""
                obj_data = conn.find('DTS:ObjectData', self.ns)
                if obj_data is not None:
                    inner_conn = obj_data.find('DTS:ConnectionManager', self.ns)
                    if inner_conn is not None:
                        conn_str = inner_conn.get(f'{{{self.ns["DTS"]}}}ConnectionString', "")
                
                self.connections.append({
                    'Name': name,
                    'Type': creation_name,
                    'ConnectionString': conn_str
                })
                
                # Check for flat file columns if it's a FLATFILE connection
                if creation_name == "FLATFILE":
                    self._parse_flat_file_columns(conn, name)

    def _parse_flat_file_columns(self, conn_node, conn_name):
        obj_data = conn_node.find('DTS:ObjectData', self.ns)
        if obj_data is not None:
            inner_conn = obj_data.find('DTS:ConnectionManager', self.ns)
            if inner_conn is not None:
                columns_node = inner_conn.find('DTS:FlatFileColumns', self.ns)
                if columns_node is not None:
                    cols = []
                    for col in columns_node.findall('DTS:FlatFileColumn', self.ns):
                        # Added type mapping for readability
                        type_code = col.get(f"{{{self.ns['DTS']}}}DataType", "Unknown")
                        cols.append({
                            'Name': col.get(f'{{{self.ns["DTS"]}}}ObjectName'),
                            'DataType': SSIS_TYPE_MAP.get(type_code, f"Code {type_code}"),
                            'Width': col.get(f'{{{self.ns["DTS"]}}}MaximumWidth', 'N/A')
                        })
                    self.flat_file_columns.append({
                        'ConnectionName': conn_name,
                        'Columns': cols
                    })

    def _parse_variables(self):
        variables_node = self.root.find('DTS:Variables', self.ns)
        if variables_node is not None:
            for var in variables_node.findall('DTS:Variable', self.ns):
                name = var.get(f'{{{self.ns["DTS"]}}}ObjectName')
                # Variables often have a Namespace as well
                ns_prefix = var.get(f'{{{self.ns["DTS"]}}}Namespace', "User")
                full_name = f"{ns_prefix}::{name}"
                
                val_node = var.find('DTS:VariableValue', self.ns)
                data_type = val_node.get(f'{{{self.ns["DTS"]}}}DataType') if val_node is not None else "Unknown"
                value = val_node.text if val_node is not None else ""
                
                self.variables.append({
                    'Name': full_name,
                    'DataType': data_type,
                    'Value': value
                })

    def _parse_sql_and_flat_files(self):
        # We need to look inside Executables -> Pipeline -> components
        executables = self.root.find('DTS:Executables', self.ns)
        if executables is not None:
            self._scan_executables(executables)

    def _scan_executables(self, parent_node):
        for exec_node in parent_node.findall('DTS:Executable', self.ns):
            # Check for Pipeline tasks
            obj_data = exec_node.find('DTS:ObjectData', self.ns)
            if obj_data is not None:
                pipeline = obj_data.find('pipeline')
                if pipeline is not None:
                    components = pipeline.find('components')
                    if components is not None:
                        for comp in components.findall('component'):
                            name = comp.get('name')
                            props = comp.find('properties')
                            if props is not None:
                                for prop in props.findall('property'):
                                    prop_name = prop.get('name')
                                    if prop_name in ['SqlCommand', 'SqlCommandParam']:
                                        if prop.text and prop.text.strip():
                                            self.sql_queries.append({
                                                'ComponentName': name,
                                                'SQL': prop.text.strip()
                                            })
            
            # Recurse into child executables (e.g. Foreach Loop, Sequence Container)
            child_execs = exec_node.find('DTS:Executables', self.ns)
            if child_execs is not None:
                self._scan_executables(child_execs)

def generate_word_doc(parser, output_path):
    doc = Document()
    doc.add_heading('DTSX Package Specification', 0)

    # File Info
    doc.add_heading('Package Information', level=1)
    doc.add_paragraph(f'Source File: {os.path.basename(parser.file_path)}')

    # Connections
    doc.add_heading('Connection Managers', level=1)
    if parser.connections:
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Type'
        hdr_cells[2].text = 'Connection String'
        for conn in parser.connections:
            row_cells = table.add_row().cells
            row_cells[0].text = conn['Name']
            row_cells[1].text = conn['Type']
            row_cells[2].text = conn['ConnectionString']
    else:
        doc.add_paragraph('No connection managers found.')

    # Variables
    doc.add_heading('Variables', level=1)
    if parser.variables:
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Data Type'
        hdr_cells[2].text = 'Value'
        for var in parser.variables:
            row_cells = table.add_row().cells
            row_cells[0].text = var['Name']
            row_cells[1].text = var['DataType']
            row_cells[2].text = str(var['Value'])
    else:
        doc.add_paragraph('No variables found.')

    # SQL Code
    doc.add_heading('Extracted SQL Code', level=1)
    if parser.sql_queries:
        for sql_item in parser.sql_queries:
            doc.add_heading(f"Component: {sql_item['ComponentName']}", level=2)
            doc.add_paragraph(sql_item['SQL'])
    else:
        doc.add_paragraph('No SQL code extracted.')

    # Flat File Columns
    doc.add_heading('Flat File Column Definitions', level=1)
    if parser.flat_file_columns:
        for ff in parser.flat_file_columns:
            doc.add_heading(f"Connection: {ff['ConnectionName']}", level=2)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Column Name'
            hdr_cells[1].text = 'Data Type'
            hdr_cells[2].text = 'Width'
            for col in ff['Columns']:
                row_cells = table.add_row().cells
                row_cells[0].text = col['Name']
                row_cells[1].text = col['DataType']
                row_cells[2].text = str(col['Width'])
    else:
        doc.add_paragraph('No flat file columns found.')

    doc.save(output_path)
    print(f"Spec document generated at: {output_path}")

if __name__ == "__main__":
    dtsx_path = "Lesson 1.dtsx"
    output_docx = "Spec_Document.docx"
    
    if os.path.exists(dtsx_path):
        parser = DTSXParser(dtsx_path)
        parser.parse()
        generate_word_doc(parser, output_docx)
    else:
        print(f"Error: {dtsx_path} not found.")
